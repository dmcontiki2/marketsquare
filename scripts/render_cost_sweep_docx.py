#!/usr/bin/env python3
"""
render_cost_sweep_docx.py — render a COST_SWEEP_<date>.md report into David's
styled .docx (Heading 0 title, Heading 1 sections, severity-coloured bullets,
bold totals line). Pure formatting, no AI, no network.

Usage: python3 render_cost_sweep_docx.py <in.md> <out.docx>
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

SEVERITY_COLORS = {
    "CRITICAL": RGBColor(0xC0, 0x28, 0x28),
    "WARN":     RGBColor(0xB4, 0x6A, 0x00),
    "OK":       RGBColor(0x1E, 0x7D, 0x32),
    "INFO":     RGBColor(0x44, 0x55, 0x66),
}

# Emoji badges used by cost_compliance_sweep.py's render(): CRITICAL=red circle,
# WARN=orange circle, INFO=info symbol, OK=check mark box. Stripped because the
# severity word is already bolded + coloured immediately after. Plain check marks
# (U+2713, used inline e.g. "ceiling [check] spend-log [check]") are kept — they
# aren't pictographic emoji and carry real information.
EMOJI_RE = re.compile(
    "[✅ℹ️\U0001F534\U0001F7E0⚠❌]"
)
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def clean(text: str) -> str:
    text = EMOJI_RE.sub("", text)
    text = text.replace("`", "")
    return re.sub(r"\s+", " ", text).strip()


def add_bold_runs(paragraph, content: str, force_bold: bool = False):
    """Split content on **bold** markdown, adding runs to an existing paragraph.
    Bold segments matching a known severity word get that severity's colour."""
    pos = 0
    for m in BOLD_RE.finditer(content):
        if m.start() > pos:
            r = paragraph.add_run(content[pos:m.start()])
            if force_bold:
                r.bold = True
        word = m.group(1)
        r = paragraph.add_run(word)
        r.bold = True
        if word in SEVERITY_COLORS:
            r.font.color.rgb = SEVERITY_COLORS[word]
        pos = m.end()
    if pos < len(content):
        r = paragraph.add_run(content[pos:])
        if force_bold:
            r.bold = True


def build(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    i = 0

    if lines and lines[0].startswith("# "):
        doc.add_heading(clean(lines[0][2:]), level=0)
        i = 1

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(clean(stripped[4:]), level=2)
        elif stripped.startswith("## "):
            doc.add_heading(clean(stripped[3:]), level=1)
        elif stripped.startswith("# "):
            doc.add_heading(clean(stripped[2:]), level=0)
        elif stripped.startswith("- "):
            content = clean(stripped[2:])
            p = doc.add_paragraph(style="List Bullet")
            add_bold_runs(p, content)
        elif stripped.startswith("**Totals:**"):
            content = clean(stripped)
            p = doc.add_paragraph()
            add_bold_runs(p, content, force_bold=True)
        elif stripped.startswith("_") and stripped.endswith("_") and len(stripped) > 1:
            p = doc.add_paragraph()
            r = p.add_run(clean(stripped[1:-1]))
            r.italic = True
        else:
            doc.add_paragraph(clean(stripped))
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: render_cost_sweep_docx.py <in.md> <out.docx>", file=sys.stderr)
        sys.exit(1)
    build(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"docx -> {sys.argv[2]}")
